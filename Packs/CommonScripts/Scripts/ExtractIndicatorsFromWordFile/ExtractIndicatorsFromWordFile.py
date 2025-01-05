import demistomock as demisto
from CommonServerPython import *
from CommonServerUserPython import *
import subprocess
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.exceptions import PackageNotFoundError


class WordParser:

    def __init__(self):
        self.res = []   # type: List[Dict[str, str]]
        self.errEntry = {
            "Type": entryTypes["error"],
            "ContentsFormat": formats["text"],
            "Contents": ""
        }
        self.file_path = ""
        self.file_name = ""
        self.file_type = ""
        self.paragraphs = ""
        self.tables = ""
        self.hyperlinks = ""
        self.core_properties = {}

    def get_file_details(self):
        file_path_data = demisto.getFilePath(demisto.args().get("entryID"))
        self.file_path = file_path_data.get('path')
        self.file_name = file_path_data.get('name')
        file_entry = demisto.dt(self.get_context(), "File(val.EntryID === '{}')".format(demisto.args().get('entryID')))
        if isinstance(file_entry, list):
            file_entry = file_entry[0]

        self.file_type = file_entry.get("Type")

    def convert_doc_to_docx(self):
        output = subprocess.check_output(
            ['soffice', '--headless', '-env:UserInstallation=file:///tmp/.config/extractindicators', '--convert-to',
             'docx', self.file_path], stderr=subprocess.STDOUT)
        demisto.debug(f"soffice output: [{str(output)}]")
        # Requires office-utils docker image
        output_file_name = self.file_name[0:self.file_name.rfind('.')] + '.docx'
        self.file_path = self.file_path + ".docx"
        try:
            with open(self.file_path, 'rb') as f:
                f_data = f.read()
                self.res = fileResult(output_file_name, f_data)
        except OSError:
            return_error("Error: was not able to convert the input file to docx format.")

    def extract_indicators(self):
        try:
            document = Document(self.file_path)
            self.paragraphs = self.get_paragraphs(document)
            self.tables = self.get_tables(document)
            self.hyperlinks = self.get_hyperlinks(document)
            self.get_core_properties(document)
        except PackageNotFoundError:
            self.errEntry["Contents"] = "Input file is not a valid docx/doc file."
            self.res = self.errEntry  # type: ignore
        except BaseException as e:
            self.errEntry["Contents"] = "Error occurred while parsing input file.\nException info: " + str(e)
            self.res = self.errEntry  # type: ignore

    def get_paragraphs(self, document):
        return '\n'.join([para.text for para in document.paragraphs])

    def get_tables(self, document):
        all_cells_txt = ""
        if document.tables:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_cells_txt += (" " + para.text)
        return " ".join(all_cells_txt.split())  # Removes extra whitespaces

    def get_core_properties(self, document):
        self.core_properties['author'] = document.core_properties.author
        self.core_properties['category'] = document.core_properties.category
        self.core_properties['comments'] = document.core_properties.comments
        self.core_properties['identifier'] = document.core_properties.identifier
        self.core_properties['keywords'] = document.core_properties.keywords
        self.core_properties['subject'] = document.core_properties.subject
        self.core_properties['title'] = document.core_properties.title

    def get_hyperlinks(self, document):
        all_hyperlinks = ""
        rels = document.part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                all_hyperlinks += (rels[rel]._target + " ")
        return all_hyperlinks

    def parse_word(self):
        self.get_file_details()
        if self.file_name.lower().endswith(".doc") or "Composite Document File V2 Document" in self.file_type:
            self.convert_doc_to_docx()
            self.extract_indicators()
        elif self.file_name.lower().endswith(".docx") or "Microsoft Word 2007+" in self.file_type:
            self.extract_indicators()
        else:
            return_error("Input file is not a doc file.")

    def get_context(self):
        incident_id = demisto.incident()['id']
        res = execute_command('getContext', {'id': incident_id})
        return demisto.get(res, 'context')


def main():
    # Parsing:
    parser = WordParser()
    try:
        parser.parse_word()
    except subprocess.CalledProcessError as perr:
        return_error(f"ProcessError: exit code: {perr.returncode}. Output: {perr.output}")
    except Exception as e:
        return_error(str(e))

    core_properties_str = " ".join(prop for prop in parser.core_properties.values())
    all_data = " ".join([parser.paragraphs, parser.tables, parser.hyperlinks, core_properties_str])

    # Returning Indicators:
    indicators_hr = demisto.executeCommand("extractIndicators", {'text': all_data})
    return_results(indicators_hr)

    hr_output = tableToMarkdown('Properties', parser.core_properties)
    hr_output += f'### Paragraphs\n{parser.paragraphs}\n'
    hr_output += f'### Tables\n{parser.tables}\n'
    hr_output += f'### Hyperlinks\n{parser.hyperlinks}'

    # Returning all parsed data:
    return_results(CommandResults(readable_output=hr_output))

    # Returning error:
    if parser.res:  # If there was an error:
        contents = parser.res["Contents"]  # type: ignore
        if "Error occurred while parsing input" in contents or "Input file is not a valid" in contents:
            return_results(parser.res)  # Return error too


# python2 uses __builtin__ python3 uses builtins
if __name__ == "__builtin__" or __name__ == "builtins":
    main()
