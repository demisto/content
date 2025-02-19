import pytest
import os
import shutil
from ExtractIndicatorsFromWordFile import WordParser, main
import demistomock as demisto
from unittest.mock import MagicMock

expected_partial_all_data = 'Lorem ipsum dolor sit amet, an quas nostro posidonium mei, pro choro vocent pericula et'


@pytest.mark.parametrize('file_name,file_path', [
    ('docwithindicators.doc', 'test_data/docwithindicators'),
    ('docxwithindicators.docx', 'test_data/docxwithindicators')
])
def test_parse_word(file_name, file_path, request):
    basename = os.path.basename(file_path)
    shutil.copy(file_path, os.getcwd())

    if os.getcwd().endswith('test_data'):
        os.chdir('..')
    parser = WordParser()
    parser.get_file_details = lambda: None
    parser.file_name = file_name
    parser.file_path = basename
    parser.parse_word()
    assert (expected_partial_all_data in parser.paragraphs)


def test_getting_file_from_context(mocker):
    """
    Given:
        -

    When:
        - Call to get_file_details

    Then:
        - Validate the context comes from incident by calling execute_command('getContext')
        instead of demisto.context witch is missed in context of sub-playbook

    """

    # prepare
    parser = WordParser()
    mocker.patch.object(demisto, 'dt')
    mocker.patch.object(demisto, 'args', return_value={})
    mocker.patch.object(demisto, 'incident', return_value={'id': 1})
    mocked_method = mocker.patch('ExtractIndicatorsFromWordFile.execute_command', return_value={'context': {}})

    # run
    parser.get_file_details()

    # validate
    assert mocked_method.call_args[0][0] == 'getContext'


def test_get_hyperlinks():
    """
    Given:
        - Document with 3 hyperlinks.
    When:
        - Call get_hyperlinks
    Then:
        - Validate the result contains the 3 links.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    parser = WordParser()

    doc = MagicMock()
    doc.part.rels = {}

    doc.part.rels = {
        "r1": MagicMock(reltype=RT.HYPERLINK, _target="http://example1.com"),
        "r2": MagicMock(reltype=RT.HYPERLINK, _target="http://example2.com"),
        "r3": MagicMock(reltype=RT.HYPERLINK, _target="http://example3.com")
    }

    result = parser.get_hyperlinks(doc)
    assert result == "http://example1.com http://example2.com http://example3.com "


def test_get_paragraphs():
    """
    Given:
        - Document with 3 paragraphs.
    When:
        - Call get_paragraphs
    Then:
        - Validate the result contains the 3 paragraphs.
    """
    parser = WordParser()

    doc = MagicMock()
    mock_paragraphs = [
        MagicMock(text='This is the first paragraph.'),
        MagicMock(text='This is the second paragraph.'),
        MagicMock(text='This is the third paragraph.')
    ]

    doc.paragraphs = mock_paragraphs
    result = parser.get_paragraphs(doc)
    assert result == "This is the first paragraph.\nThis is the second paragraph.\nThis is the third paragraph."


def test_get_tables():
    """
    Given:
        - Document with a table.
    When:
        - Call get_tables
    Then:
        - Validate the result contains the table data.
    """
    parser = WordParser()

    doc = MagicMock()

    # create a mock table, row, cell, and paragraph structure
    mock_table = MagicMock()
    mock_row = MagicMock()
    mock_cell_1 = MagicMock()
    mock_cell_2 = MagicMock()

    mock_paragraph_1 = MagicMock(text="Cell 1 text.")
    mock_paragraph_2 = MagicMock(text="Cell 2 text.")

    mock_cell_1.paragraphs = [mock_paragraph_1]
    mock_cell_2.paragraphs = [mock_paragraph_2]

    mock_row.cells = [mock_cell_1, mock_cell_2]
    mock_table.rows = [mock_row]
    doc.tables = [mock_table]

    result = parser.get_tables(doc)
    assert result == "Cell 1 text. Cell 2 text."


def test_main(mocker):
    """
    Given:
        - Document with author, title and paragraph.
    When:
        - Call main
    Then:
        - Validate the human-readable output is correct.
        - validate extractIndicators command executed with the correct arguments.
    """

    expected_hr_output = '''### Properties
    |author|title|
    |---|---|
    | author | title |

    ### Paragraphs
    paragraphs

    ### Tables

    ### Hyperlinks
    '''

    parser = WordParser()
    parser.core_properties = {'author': 'author', 'title': 'title'}
    parser.paragraphs = 'paragraphs'
    mocker.patch('ExtractIndicatorsFromWordFile.WordParser', return_value=parser)

    mocker.patch.object(parser, 'parse_word', return_value={})
    execute_command_mock = mocker.patch.object(demisto, 'executeCommand', return_value={})
    return_results_mock = mocker.patch('ExtractIndicatorsFromWordFile.return_results')

    main()

    return_results_mock.call_args[0][0] == expected_hr_output
    assert execute_command_mock.call_args[0][0] == 'extractIndicators'
    assert execute_command_mock.call_args[0][1]['text'] == 'paragraphs   author title'
