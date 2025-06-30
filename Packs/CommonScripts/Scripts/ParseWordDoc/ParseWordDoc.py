import zipfile
from xml.etree.ElementTree import XML

import demistomock as demisto  # noqa: F401
from CommonServerPython import *  # noqa: F401
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.exceptions import PackageNotFoundError
from docx.document import Document as DocumentObject

import subprocess
from tempfile import TemporaryDirectory
from pathlib import Path

def convert_to_docx(doc_path, output_dir):
    """Convert a document to ".docx". Used for old files that can't be parsed conventionally."""
    command = [
        'soffice',
        '--headless',
        '--convert-to', 'docx',
        '--outdir', output_dir,
        doc_path
    ]

    result = subprocess.run(command, capture_output=True, text=True)

    demisto.debug(f"LibreOffice stdout: {result.stdout.strip()!r}, stderr: {result.stderr.strip()}")


def get_document_urls(file_path: str) -> tuple[DocumentObject, list[str]]:
    try:
        document = Document(file_path)
        return document, extract_urls_xml(file_path) + extract_urls_docx(document)
    except Exception as e:
        if "themeManager+xm1" in str(e):
            demisto.debug(f'Unable to parse doc: "{e}". Retrying with LibreOffice')
            with TemporaryDirectory() as d:
                convert_to_docx(file_path, d)
                file_path = str(Path(d) / replace_suffix(file_path, 'docx'))
                document = Document(file_path)
                return document, extract_urls_xml(file_path) + extract_urls_docx(document)
        else:
            raise


def extract_urls_xml(file_path: str) -> list[str]:
    urls = []
    document = zipfile.ZipFile(file_path)
    xml_content = document.read("word/document.xml")
    document.close()
    tree = XML(xml_content)

    for element in tree.iter():
        if hasattr(element, "text") and element.text and "HYPERLINK" in element.text:
            url = element.text.replace(' HYPERLINK "', "")[:-1]
            urls.append(url)
    return urls


def extract_urls_docx(document: DocumentObject) -> list[str]:
    urls = []
    rels = document.part.rels
    for rel in rels.values():
        if rel.reltype == RT.HYPERLINK:
            urls.append(rel._target)
    return urls


def replace_suffix(file_name: str, new_suffix: str) -> str:
    return f"{file_name.partition('.')[0]}.{new_suffix}"


def parse_word_doc(entry_id):

    try:
        cmd_res = demisto.getFilePath(entry_id)
        file_path = cast(str, cmd_res.get("path"))
        demisto.debug(f'{file_path=}')
        document, urls = get_document_urls(file_path)
    except PackageNotFoundError:
        return_error("Input file is not a valid docx/doc file.")
    except BaseException as e:
        return_error(f"Error occurred while parsing input file.\nException info: {e}")

    file_data = (
        "\n".join(para.text for para in document.paragraphs)
        + "\n\n\nExtracted links:\n* "
        + "\n* ".join(urls))
    file_name = replace_suffix(cmd_res['name'], "txt")
    res = fileResult(file_name, file_data.encode("utf8"))

    demisto.results(res)


def main():
    entry_id = demisto.args()["entryID"]
    parse_word_doc(entry_id)


if __name__ in ("__main__", "__builtin__", "builtins"):
    main()
