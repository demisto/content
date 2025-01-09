import demistomock as demisto  # noqa: F401
from CommonServerPython import *  # noqa: F401
# import json
from docx import Document
from docx.enum.section import WD_ORIENT   # pylint: disable=E0611
from docx.shared import Inches

# Purpose:      This automation will produce docx file detailing the tasks in the given playbook. It can
#               produce a table or paragraph format of the report.
#               The automation will create seprate tables (in case of Table format), or separate paragraphs
#               (in case of Paragraph format).
#               The definition of a section is path starting from a section header type of task, and ends
#               when it hits another section header in that path pf traversal.
# Author:       Mahmood Azmat
# Input1:       Name of the playbook (Mandatory)
# Input2:       Format type needed. Table or Paragraph. Paragraph is default.
# Input3:       Name of the docx file that will be produced. Give the full name including the ".docx"
#               extension. (Mandatory)
# Requirements: This automation requires "Core REST API" integration enabled and connected to the XSOAR itself.
#               Automation uses it to read the objects of the playbook.

retVal: dict = {}
DEMISTO_PLAYBOOKS_PATH = "/playbook/search"
document = Document()
SectionID = ""
Table = ""
Paragraph = ""
TaskIDsInLogic = []
TaskIDsInLogic.append('0')
OutPutType = ""
ConditionBranch = ""
PlaybookIndex = -1
NodesVisitedInThisSection = []
NodesVisitedInAllSections = {}
#  This is a 2 level associative array. First level is the ID of the section task and
#  second level has the IDs visited in that section


def post_api_request(url, body):

    api_args = {
        "uri": url,
        "body": body
    }

    raw_res = demisto.executeCommand("core-api-post", api_args)
    try:
        res = raw_res[0]['Contents']['response']
        return res
    except KeyError:
        return_error(f'API Request failed, no response from API call to {url}')
    except TypeError:
        return_error(f'API Request failed, failedto {raw_res}')


def EnterHeader(HeaderStr):

    document.add_heading(HeaderStr, level=1)


def StartParagraph(Name, Description):

    global Paragraph
    document.add_heading(Name, level=2)

    if (Description == "[Blank]"):
        Description = ""

    Paragraph = document.add_paragraph(Description)  # type: ignore[assignment]
    Paragraph.paragraph_format.left_indent = Inches(0.25)   # type: ignore


def StartTable():

    global Table
    Table = document.add_table(rows=1, cols=2)
    Table.style = 'Light Grid Accent 1'   # type: ignore
    hdr_cells = Table.rows[0].cells   # type: ignore
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Description'


def TraverseTasks(TaskID):

    global retVal, document, Table, Paragraph, NodesVisitedInAllSections, PlaybookIndex
    global OutPutType, ConditionBranch, NodesVisitedInThisSection, SectionID

    if (TaskID in NodesVisitedInThisSection):

        # If this node has already been visited then we will skip it. This avoids putting the same task into the
        # document mutiple times.
        # You may have mutiple paths ending to the same task within the same section so
        # we do not want to make multiple entries of that.
        return

    else:

        NodesVisitedInThisSection.append(TaskID)
        NodesVisitedInAllSections[SectionID] = NodesVisitedInThisSection

    if (TaskID is None):
        return

    else:

        curTask = retVal[PlaybookIndex]['tasks'][str(TaskID)]

        taskDescription = curTask["task"].get("description", "[Blank]")

        Description = taskDescription
        Type = curTask['type']

        if (str(TaskID) == '0'):
            Name = 'Playbook Triggered'

        else:
            Name = curTask['task']['name']

            if (Type == 'condition'):
                Name = Name + " - Conditional Task "

            if (Type == 'playbook'):
                Name = Name + " - Sub-playbook "

            if (ConditionBranch != ""):
                Name = Name + " (Condition branch->" + ConditionBranch + ")"

        if (Type == 'title' or Type == 'start'):

            if (Type == 'title'):
                if (curTask.get('nextTasks') is not None):
                    EnterHeader(Name)
                    if (OutPutType == 'TABLE'):
                        StartTable()

            else:
                EnterHeader(Name)
                if (OutPutType == 'TABLE'):
                    StartTable()
        else:

            if (OutPutType == 'TABLE'):
                if (Table == ""):
                    StartTable()

                row_cells = Table.add_row().cells   # type: ignore
                row_cells[0].text = Name
                row_cells[1].text = Description

            else:
                StartParagraph(Name, Description)

        if (curTask.get('nextTasks') is not None):

            if (curTask['type'] != 'condition'):
                for NextTaskID in curTask['nextTasks']['#none#']:
                    TypeOfNextTask = retVal[PlaybookIndex]['tasks'][str(NextTaskID)]['task']['type']
                    NameOfNextTask = retVal[PlaybookIndex]['tasks'][str(NextTaskID)]['task']['name']

                    if (TypeOfNextTask == 'title'):
                        NameOfNextTask = NameOfNextTask + " - Section "

                    if (ConditionBranch != ""):
                        NameOfNextTask = NameOfNextTask + " (Condition branch->" + ConditionBranch + ")"

                    if (TypeOfNextTask != 'title'):

                        TraverseTasks(NextTaskID)

                    else:
                        if (OutPutType == 'TABLE'):
                            row_cells = Table.add_row().cells   # type: ignore
                            row_cells[0].text = NameOfNextTask
                        else:
                            StartParagraph(NameOfNextTask, "")

                        if (str(NextTaskID) not in TaskIDsInLogic):
                            TaskIDsInLogic.append(str(NextTaskID))
            else:
                for TempKeys in curTask.get('nextTasks'):

                    for NextTaskID in curTask['nextTasks'][TempKeys]:
                        TypeOfNextTask = retVal[PlaybookIndex]['tasks'][str(NextTaskID)]['task']['type']
                        NameOfNextTask = retVal[PlaybookIndex]['tasks'][str(NextTaskID)]['task']['name']

                        ConditionBranch = TempKeys

                        if (TypeOfNextTask == 'title'):
                            NameOfNextTask = NameOfNextTask + " - Section "

                        if (ConditionBranch == "#default#"):
                            ConditionBranch = "Else"

                        NameOfNextTask = NameOfNextTask + " (Condition branch->" + ConditionBranch + ")"

                        if (TypeOfNextTask != 'title'):

                            TraverseTasks(NextTaskID)

                        else:
                            if (OutPutType == 'TABLE'):
                                row_cells = Table.add_row().cells   # type: ignore
                                row_cells[0].text = NameOfNextTask
                            else:
                                StartParagraph(NameOfNextTask, "")

                            if (str(NextTaskID) not in TaskIDsInLogic):
                                TaskIDsInLogic.append(str(NextTaskID))

                ConditionBranch = ""

    return


''' MAIN FUNCTION '''


def main():

    # print ("Test")
    global retVal, TaskIDsInLogic, NodesVisitedInThisSection, PlaybookIndex
    global document, SectionID
    global OutPutType
    global ConditionBranch
    args = demisto.args()
    DocFileName = args.get('DocFileName')
    PlaybookName = args.get('PlaybookName')
    OutPutType = args.get('Output_Format')
    OutPutType = str(OutPutType).upper()

    if (OutPutType == ""):
        OutPutType = "PARAGRAPH"

    TempStr = PlaybookName + "- Playbook"

    if (OutPutType == 'TABLE'):
        sections = document.sections
        section = sections[-1]

        # Swapping width with height to make the table landscape format
        old_width = section.page_width
        old_height = section.page_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = old_height
        section.page_height = old_width

    document.add_heading(TempStr, 0)
    retVal = post_api_request(DEMISTO_PLAYBOOKS_PATH, {"query": PlaybookName}).get("playbooks")

    if (retVal is None):
        demisto.results("Playbook not found. Please check the Playbook name.")
        return

    i = 0
    j = 0

    while j < len(retVal):

        PlaybookNameTempStr = retVal[j].get('nameRaw')

        if (PlaybookNameTempStr == PlaybookName):
            PlaybookIndex = j

        j = j + 1

    if (PlaybookIndex == -1):
        demisto.results("Playbook not found. Please check the Playbook name.")
        return

    while (TaskIDsInLogic[i] is not None):
        ConditionBranch = ""
        NodesVisitedInThisSection.clear()
        SectionID = TaskIDsInLogic[i]
        TraverseTasks(str(TaskIDsInLogic[i]))
        i = i + 1

        try:
            assert (TaskIDsInLogic[i])

        except Exception:
            break

    if (OutPutType == 'TABLE'):
        for tb in document.tables:
            for cell in tb.columns[0].cells:
                cell.width = Inches(3)

            for cell in tb.columns[1].cells:
                cell.width = Inches(6)

    document.save(DocFileName)

    with open(DocFileName, 'rb') as f:
        filedata = f.read()

    demisto.results(fileResult(DocFileName, filedata))
    return


''' ENTRY POINT '''

if __name__ in ('__main__', '__builtin__', 'builtins'):
    main()
